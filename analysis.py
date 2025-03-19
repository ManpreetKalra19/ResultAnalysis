import streamlit as st
import pandas as pd
import numpy as np
import io
import base64
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import seaborn as sns

# Set page configuration
st.set_page_config(
    page_title="Class XII Academic Result Analysis",
    page_icon="📊",
    layout="wide"
)

# Helper functions for document creation
def set_cell_border(cell, border_size=8):
    """Set cell border size."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create border elements
    for border_position in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'w:tblBorders')
        border_element = OxmlElement(f'w:{border_position}')
        border_element.set(qn('w:val'), 'single')
        border_element.set(qn('w:sz'), str(border_size))
        border_element.set(qn('w:space'), '0')
        border_element.set(qn('w:color'), 'auto')
        border.append(border_element)
        tcPr.append(border)

def add_heading(doc, text, level=1):
    """Add heading with underline for level 1."""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        for run in heading.runs:
            run.underline = True
    return heading

def create_table(doc, rows, cols, data=None):
    """Create table with borders."""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    # If data is provided, fill the table
    if data is not None:
        for i, row_data in enumerate(data):
            for j, cell_value in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = str(cell_value)
                set_cell_border(cell)
    
    return table

def generate_download_link(doc, filename):
    """Generate a download link for the document."""
    # Save document to a BytesIO object
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Create download link
    b64 = base64.b64encode(file_stream.read()).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download {filename}</a>'
    return href

# Main app functions
def load_data(result_file, subject_codes_file):
    """Load and process the uploaded CSV files."""
    try:
        # Load results data
        results_df = pd.read_csv(result_file)
        
        # Load subject codes data
        subject_codes_df = pd.read_csv(subject_codes_file, skiprows=0)
        #subject_codes_df = pd.read_csv(subject_codes_file, skiprows=1)

        
        # Clean up subject codes dataframe
        subject_codes_df.columns = [col.strip() for col in subject_codes_df.columns]
        subject_codes_df = subject_codes_df.dropna(axis=1, how='all')
        
        # Extract columns with data
        valid_columns = []
        for col in subject_codes_df.columns:
            if not pd.isna(subject_codes_df[col]).all():
                valid_columns.append(col)
        
        subject_codes_df = subject_codes_df[valid_columns]
        subject_codes_df.columns = ['Subject', 'Code']
        
        # Clean up subject codes data
        subject_codes_df['Subject'] = subject_codes_df['Subject'].apply(lambda x: str(x).strip())
        #subject_codes_df['Code'] = subject_codes_df['Subject'].apply(lambda x: str(x).split('\n')[-1].strip() if '\n' in str(x) else '')
        subject_codes_df['Subject'] = subject_codes_df['Subject'].apply(lambda x: str(x).split('\n')[0].strip() if '\n' in str(x) else x)
        
        return results_df, subject_codes_df
    
    except Exception as e:
        st.error(f"Error loading data: {e}")
        return None, None
    

def analyze_results(results_df):
    """Analyze the results data and return metrics."""
    try:
        # Basic metrics
        total_students = len(results_df)
        
        # Extract all subject marks columns
        subject_columns = []
        for i in range(0,6):  # Assuming maximum 6 subjects per student
            code_col = f'Code.{i}' if i > 0 else 'Code'
            marks_col = f'Marks.{i}' if i > 0 else 'Marks'
            
            if code_col in results_df.columns and marks_col in results_df.columns:
                subject_columns.append((code_col, marks_col))
            else:
                # Try without the dot notation for first columns
                if i == 0:
                    if 'Code' in results_df.columns and 'Marks' in results_df.columns:
                        subject_columns.append(('Code', 'Marks'))
        
        # Calculate metrics for each student
        students_data = []
        compartment_count = 0
        failed_count = 0
        
        for _, row in results_df.iterrows():
            student_marks = []
            subject_marks_dict = {}  # To store subject code and marks pairs
            
            for code_col, marks_col in subject_columns:
                if pd.notna(row[code_col]) and pd.notna(row[marks_col]) and row[marks_col] > 0:
                    student_marks.append(row[marks_col])
                    subject_marks_dict[str(row[code_col])] = row[marks_col]
            
            if student_marks:
                # Calculate best of five
                if len(student_marks) >= 5:
                    # Sort marks in descending order
                    sorted_marks = sorted(student_marks, reverse=True)
                    best_five = sorted_marks[:5]
                    avg_best_five = sum(best_five) / 5
                    
                    # Check for compartment and failed status in best five subjects
                    best_five_subject_codes = []
                    for code, marks in subject_marks_dict.items():
                        if marks in best_five:
                            best_five_subject_codes.append((code, marks))
                            if len(best_five_subject_codes) == 5:
                                break
                    
                    # Count subjects with marks less than 33 in best five
                    below_passing_count = sum(1 for _, marks in best_five_subject_codes if marks < 33)
                    
                    # Update compartment and failed counts
                    if below_passing_count == 1:
                        compartment_count += 1
                        status = "Compartment"
                    elif below_passing_count > 1:
                        failed_count += 1
                        status = "Failed"
                    else:
                        status = "Passed"
                else:
                    avg_best_five = sum(student_marks) / len(student_marks)
                    # For students with less than 5 subjects, check all subjects
                    below_passing_count = sum(1 for marks in student_marks if marks < 33)
                    if below_passing_count == 1:
                        compartment_count += 1
                        status = "Compartment"
                    elif below_passing_count > 1:
                        failed_count += 1
                        status = "Failed"
                    else:
                        status = "Passed"
                
                students_data.append({
                    'Name': row['Name'],
                    'Roll No': row['Roll No'],
                    'Best of Five Percentage': avg_best_five,
                    'Status': status
                })
        
        students_df = pd.DataFrame(students_data)
        
        # Calculate passed students count
        passed_count = total_students - compartment_count - failed_count
        
        # Overall school average
        school_avg = students_df['Best of Five Percentage'].mean()
        
        # Distribution of marks
        ranges = [
            (95, 101, '> 95'),
            (90, 95, '90-95'),
            (85, 90, '85-90'),
            (80, 85, '80-85'),
            (75, 80, '75-80'),
            (70, 75, '70-75'),
            (65, 70, '65-70'),
            (0, 65, '< 65')
        ]
        
        distribution = {}
        for lower, upper, label in ranges:
            count = len(students_df[(students_df['Best of Five Percentage'] >= lower) & 
                                   (students_df['Best of Five Percentage'] < upper)])
            distribution[label] = count
        
        # Calculate top 5 students
        top_students = students_df.sort_values(by='Best of Five Percentage', ascending=False).head(5)
        
        # UPDATED SECTION: Subject-wise analysis using the new approach
        subject_data = []
        for i in range(6):  # Assuming maximum 6 subjects per student
            if i == 0:
                code_col = 'Code'
                marks_col = 'Marks'
            else:
                code_col = f'Code.{i}'
                marks_col = f'Marks.{i}'
                
            if code_col in results_df.columns and marks_col in results_df.columns:
                temp = results_df[['Name', code_col, marks_col]].copy()
                temp = temp.rename(columns={code_col: 'Subject', marks_col: 'Marks'})
                # Drop rows where Subject is NaN or 0
                temp = temp.dropna(subset=['Subject'])
                temp = temp[temp['Subject'] != 0]
                temp = temp[temp['Subject'] != '0']
                # Convert marks to numeric
                temp['Marks'] = pd.to_numeric(temp['Marks'], errors='coerce')
                # Only include rows with valid marks
                temp = temp[temp['Marks'] > 0]
                subject_data.append(temp)
        
        # Concatenate all subject dataframes
        if subject_data:
            long_df = pd.concat(subject_data, ignore_index=True)
            
            # Group by Subject and aggregate
            subject_analysis = {}
            groups = long_df.groupby('Subject')
            
            for subject, group in groups:
                if str(subject) != '0':  # Skip subject code 0
                    highest_mark = group['Marks'].max()
                    num_students = group['Name'].nunique()
                    subject_avg = group['Marks'].mean()
                    top_scorers = group[group['Marks'] == highest_mark]['Name'].unique()
                    
                    subject_analysis[str(subject)] = {
                        'Subject Code': str(subject),
                        'Highest Marks': highest_mark,
                        'Students': num_students,
                        'Subject Average': subject_avg,
                        'Toppers': list(top_scorers)
                    }
        else:
            subject_analysis = {}
        
        return {
            'total_students': total_students,
            'passed_count': passed_count,
            'compartment_count': compartment_count,
            'failed_count': failed_count,
            'school_avg': school_avg,
            'distribution': distribution,
            'top_students': top_students,
            'subject_analysis': subject_analysis,
            'students_df': students_df  # Include full student data for additional analyses
        }
    
    except Exception as e:
        st.error(f"Error analyzing results: {e}")
        return None


def generate_report(analysis_results, subject_codes_df):
    """Generate a Word document report based on the analysis."""
    try:
        # Create a new Document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Result Analysis AISSCE Year 12', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Student Attendance section
        add_heading(doc, 'STUDENT ATTENDANCE', level=1)
        
        # Calculate percentages
        passed_percentage = (analysis_results['passed_count'] / analysis_results['total_students']) * 100
        compartment_percentage = (analysis_results['compartment_count'] / analysis_results['total_students']) * 100
        failed_percentage = (analysis_results['failed_count'] / analysis_results['total_students']) * 100
        
        attendance_data = [
            ['', 'NUMBER', 'PERCENTAGE'],
            ['NUMBER OF STUDENTS ENROLLED', str(analysis_results['total_students']), '100%'],
            ['NUMBER OF STUDENTS ABSENT', '0', '0%'],  # Placeholder
            ['NUMBER OF STUDENTS APPEARED', str(analysis_results['total_students']), '100%'],
            ['NUMBER OF STUDENTS PASSED', str(analysis_results['passed_count']), f"{passed_percentage:.2f}%"],
            ['NUMBER OF STUDENTS GETTING COMPARTMENTS', str(analysis_results['compartment_count']), f"{compartment_percentage:.2f}%"],
            ['NUMBER OF STUDENTS FAILED', str(analysis_results['failed_count']), f"{failed_percentage:.2f}%"]
        ]
        
        attendance_table = create_table(doc, rows=len(attendance_data), cols=3, data=attendance_data)
        
        # Average Aggregate Percentage
        doc.add_paragraph()
        avg_para = doc.add_paragraph()
        avg_run = avg_para.add_run(f"Average Aggregate Percentage of the school- {analysis_results['school_avg']:.2f}% (Best of Five)")
        avg_run.underline = True
        avg_run.bold = True
        
        # Mark Distribution
        doc.add_paragraph()
        for label, count in analysis_results['distribution'].items():
            doc.add_paragraph(f"-Students have secured {label} = {count}")
        
        # Rank Holders
        doc.add_paragraph()
        rank_para = doc.add_paragraph()
        rank_run = rank_para.add_run("Rank Holders (Best of Five)")
        rank_run.underline = True
        rank_run.bold = True
        
        rank_data = [['RANK', 'NAME OF STUDENT', 'PERCENTAGE']]
        for i, (_, student) in enumerate(analysis_results['top_students'].iterrows(), 1):
            rank_data.append([
                str(i),
                student['Name'],
                f"{student['Best of Five Percentage']:.2f}%"
            ])
        
        rank_table = create_table(doc, rows=len(rank_data), cols=3, data=rank_data)
        
        # Subject-wise Toppers
        doc.add_paragraph()
        subject_para = doc.add_paragraph()
        subject_run = subject_para.add_run("SUBJECT-WISE TOPPERS")
        subject_run.underline = True
        subject_run.bold = True
        
        subject_data = [['SUBJECT', 'SUBJECT CODE', 'HIGHEST MARKS', 'NO. OF STUDENTS', 'SUBJECT AVERAGE', 'NAME OF STUDENTS']]
        
        # Match subject codes with names
        subject_code_map = dict(zip(subject_codes_df['Code'], subject_codes_df['Subject']))
        
        for code, data in analysis_results['subject_analysis'].items():
            subject_name = subject_code_map.get(code, f"Subject {code}")
            
            subject_data.append([
                subject_name,
                code,
                f"{data['Highest Marks']:.1f}",
                str(data['Students']),
                f"{data['Subject Average']:.2f}",
                ', '.join(data['Toppers'][:20])  # Limit to 20 toppers for space
            ])
        
        subject_table = create_table(doc, rows=len(subject_data), cols=6, data=subject_data)
        
        # Add compartment and fail details
        if analysis_results['compartment_count'] > 0 or analysis_results['failed_count'] > 0:
            doc.add_paragraph()
            status_para = doc.add_paragraph()
            status_run = status_para.add_run("STUDENT STATUS DETAILS")
            status_run.underline = True
            status_run.bold = True
            
            if analysis_results['compartment_count'] > 0:
                doc.add_paragraph("Students with Compartment (Failed in one subject):")
                compartment_students = analysis_results['students_df'][analysis_results['students_df']['Status'] == 'Compartment']
                for _, student in compartment_students.iterrows():
                    doc.add_paragraph(f"- {student['Name']} (Roll No: {student['Roll No']})")
            
            if analysis_results['failed_count'] > 0:
                doc.add_paragraph("Failed Students (Failed in more than one subject):")
                failed_students = analysis_results['students_df'][analysis_results['students_df']['Status'] == 'Failed']
                for _, student in failed_students.iterrows():
                    doc.add_paragraph(f"- {student['Name']} (Roll No: {student['Roll No']})")
        
        return doc
    
    except Exception as e:
        st.error(f"Error generating report: {e}")
        return None



def create_visualizations(analysis_results):
    """Create visualizations for the dashboard."""
    # Distribution of marks
    fig1, ax1 = plt.subplots(figsize=(10, 6))
    labels = list(analysis_results['distribution'].keys())
    values = list(analysis_results['distribution'].values())
    
    ax1.bar(labels, values, color='skyblue')
    ax1.set_title('Student Performance Distribution')
    ax1.set_xlabel('Percentage Range')
    ax1.set_ylabel('Number of Students')
    plt.xticks(rotation=45)
    
    # Subject averages
    fig2, ax2 = plt.subplots(figsize=(12, 8))
    subject_names = []
    subject_avgs = []
    
    for code, data in analysis_results['subject_analysis'].items():
        subject_names.append(f"{code}")
        subject_avgs.append(data['Subject Average'])
    
    # Sort by average
    sort_idx = np.argsort(subject_avgs)
    subject_names = [subject_names[i] for i in sort_idx]
    subject_avgs = [subject_avgs[i] for i in sort_idx]
    
    ax2.barh(subject_names, subject_avgs, color='lightgreen')
    ax2.set_title('Subject-wise Average Performance')
    ax2.set_xlabel('Average Marks')
    ax2.set_ylabel('Subject Code')
    
    for i, v in enumerate(subject_avgs):
        ax2.text(v + 0.5, i, f"{v:.2f}", va='center')
    
    return fig1, fig2

# App layout
# App layout
def app():
    st.title("Class XII Academic Result Analysis App")
    
    st.markdown("""
    This app analyzes academic results data and generates a comprehensive report. 
    Upload your result.csv and subject_codes.csv files to get started.
    """)
    
    # File uploads
    col1, col2 = st.columns(2)
    
    with col1:
        results_file = st.file_uploader("Upload Results CSV File", type=["csv"])
    
    with col2:
        subject_codes_file = st.file_uploader("Upload Subject Codes CSV File", type=["csv"])
    
    if results_file and subject_codes_file:
        with st.spinner("Processing data..."):
            # Load data
            results_df, subject_codes_df = load_data(results_file, subject_codes_file)
            
            if results_df is not None and subject_codes_df is not None:
                # Display loaded data
                st.subheader("Loaded Results Data")
                st.dataframe(results_df.head())
                
                st.subheader("Loaded Subject Codes Data")
                st.dataframe(subject_codes_df)
                
                # Analyze results
                analysis_results = analyze_results(results_df)
                
                if analysis_results:
                    # Display dashboard
                    st.subheader("Results Dashboard")
                    
                    # Summary metrics
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Total Students", analysis_results['total_students'])
                    with col2:
                        st.metric("School Average", f"{analysis_results['school_avg']:.2f}%")
                    with col3:
                        st.metric("Top Score", f"{analysis_results['top_students']['Best of Five Percentage'].iloc[0]:.2f}%")
                    
                    # Additional status metrics
                    st.subheader("Student Status")
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Passed", analysis_results['passed_count'],
                                 f"{(analysis_results['passed_count']/analysis_results['total_students']*100):.1f}%")
                    with col2:
                        st.metric("Compartment", analysis_results['compartment_count'],
                                 f"{(analysis_results['compartment_count']/analysis_results['total_students']*100):.1f}%")
                    with col3:
                        st.metric("Failed", analysis_results['failed_count'],
                                 f"{(analysis_results['failed_count']/analysis_results['total_students']*100):.1f}%")
                    
                    # Display performance distribution and subject averages
                    fig1, fig2 = create_visualizations(analysis_results)
                    st.pyplot(fig1)
                    st.pyplot(fig2)
                    
                    # Top students
                    st.subheader("Top 5 Students")
                    top_df = analysis_results['top_students'][['Name', 'Roll No', 'Best of Five Percentage', 'Status']]
                    top_df['Best of Five Percentage'] = top_df['Best of Five Percentage'].apply(lambda x: f"{x:.2f}%")
                    st.dataframe(top_df)
                    
                    # Subject analysis
                    st.subheader("Subject-wise Analysis")
                    
                    # Create dataframe from subject analysis dictionary
                    subject_data = []
                    for code, data in analysis_results['subject_analysis'].items():
                        subject_data.append({
                            'Subject Code': code,
                            'Highest Marks': data['Highest Marks'],
                            'Number of Students': data['Students'],
                            'Average Marks': f"{data['Subject Average']:.2f}",
                            'Toppers': ', '.join(data['Toppers'][:3])
                        })
                    
                    subject_df = pd.DataFrame(subject_data)
                    st.dataframe(subject_df)
                    
                    # Students with compartment or failed status
                    if analysis_results['compartment_count'] > 0:
                        st.subheader("Students with Compartment")
                        compartment_df = analysis_results['students_df'][analysis_results['students_df']['Status'] == 'Compartment']
                        st.dataframe(compartment_df[['Name', 'Roll No', 'Best of Five Percentage']])
                    
                    if analysis_results['failed_count'] > 0:
                        st.subheader("Failed Students")
                        failed_df = analysis_results['students_df'][analysis_results['students_df']['Status'] == 'Failed']
                        st.dataframe(failed_df[['Name', 'Roll No', 'Best of Five Percentage']])
                    
                    # Generate report
                    st.subheader("Generate Report")
                    if st.button("Generate Detailed Report"):
                        with st.spinner("Generating report..."):
                            doc = generate_report(analysis_results, subject_codes_df)
                            if doc:
                                # Generate download link
                                download_link = generate_download_link(doc, "Result_Analysis_Report.docx")
                                st.markdown(download_link, unsafe_allow_html=True)
                                st.success("Report generated successfully!")
                else:
                    st.error("Could not analyze the results data. Please check the format.")
            else:
                st.error("Could not process the uploaded files. Please check the format.")
    else:
        # Display sample data format
        st.info("Please upload both CSV files to proceed.")
        
        with st.expander("Sample Data Format"):
            st.markdown("""
            #### Results CSV Format
            ```
            Roll No,Sex,Name,Code,Marks,Grade,Code,Marks,Grade,...
            21663950,F,STUDENT NAME,301,91,A2,37,87,B1,...
            ```
            
            #### Subject Codes CSV Format
            ```
            "Year 11 (2024-2025) Subject Codes",,,
            ,Subject name & code,Number of students,,
            1,"English 301",301,,
            ```
            """)

if __name__ == "__main__":
    app()